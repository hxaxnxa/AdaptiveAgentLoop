from fastapi import APIRouter, Depends, HTTPException, status, BackgroundTasks, UploadFile, File
from sqlalchemy.orm import Session
from typing import List, Optional
from fastapi import Form, Body
from datetime import datetime, timezone, timedelta
import pypdf
import io
import docx
import uuid
import logging
import requests

from .. import crud, models, schemas, auth
from ..database import get_db
from ..agents.quiz_grader import grade_quiz
from .. import tasks
from ..core.storage import upload_file_to_storage, get_presigned_url_for_key

# --- AI Rubric Parser & Quiz Generator ---
from ..agents.evaluation_chain import llm, get_text_from_url
from pydantic import BaseModel, Field
from langchain_core.prompts import ChatPromptTemplate

def get_text_from_presigned_url(url: str) -> str:
    """PDF/text extractor from a presigned URL using requests."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        content_type = response.headers.get("content-type", "")
        if "pdf" in content_type:
            with io.BytesIO(response.content) as f:
                reader = pypdf.PdfReader(f)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif "text" in content_type:
            return response.text
        else:
            return ""
    except Exception as e:
        raise RuntimeError(f"Failed to read material file from {url}: {e}")

# --- Models for AI structured outputs ---
class ParsedCriterion(BaseModel):
    criterion: str
    max_points: int

class ParsedRubric(BaseModel):
    rubric: List[ParsedCriterion]

class AIOption(BaseModel):
    option_text: str
    is_correct: bool

class AIQuestion(BaseModel):
    question_text: str
    question_type: str
    score: int
    options: List[AIOption]
    concept_tags: List[str] = Field(description="A list of 1-3 core concepts this question is testing.")

class AIQuiz(BaseModel):
    questions: List[AIQuestion]

class OptionUpdate(BaseModel):
    is_correct: bool

logger = logging.getLogger(__name__)

router = APIRouter(
    prefix="/api/coursework",
    tags=["coursework"]
)

# ============================================================
# Helper: File Upload to MinIO
# ============================================================
def handle_file_upload(file: UploadFile, folder: str = "uploads") -> str:
    """Helper to upload a file to MinIO and return its KEY."""
    try:
        file_extension = file.filename.split(".")[-1]
        unique_filename = f"{folder}/{uuid.uuid4()}.{file_extension}"
        file_key = upload_file_to_storage(
            file_obj=file.file,
            file_name=unique_filename,
            content_type=file.content_type
        )
        return file_key
    except Exception as e:
        logger.error(f"File upload failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"File upload failed: {e}")

# ============================================================
# Create Coursework
# ============================================================
@router.post("/classrooms/{classroom_id}", response_model=schemas.CourseworkDisplay)
async def create_new_coursework(
    classroom_id: int,
    coursework: schemas.CourseworkCreate = Body(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_teacher_user)
):
    return crud.create_coursework(db=db, coursework=coursework, classroom_id=classroom_id)

# ============================================================
# Upload File (General)
# ============================================================
@router.post("/upload-file", response_model=dict)
async def upload_file(
    file: UploadFile = File(...),
    current_user: models.User = Depends(auth.get_teacher_user)
):
    file_key = handle_file_upload(file)
    return {"file_key": file_key}

# ============================================================
# AI Quiz Generation Endpoint
# ============================================================
@router.post("/generate-quiz-from-files", response_model=schemas.QuizGenerationResponse)
def generate_quiz_with_ai(
    request: schemas.QuizGenerationRequest,
    current_user: models.User = Depends(auth.get_teacher_user)
):
    ai_quiz_gen = llm.with_structured_output(AIQuiz)
    context_text = ""
    logger.info("--- [AI Quiz Gen] Starting... ---") 

    for key in request.material_file_urls:
        try:
            logger.info(f"--- [AI Quiz Gen] Extracting text from key: {key} ---")
            context_text += get_text_from_url(key) + "\n\n"
        except Exception as e:
            logger.warning(f"--- [AI Quiz Gen] Could not read material file {key}: {e} ---", exc_info=True)

    if not context_text.strip():
        logger.warning("--- [AI Quiz Gen] Context text is empty. Aborting. ---")
        raise HTTPException(status_code=400, detail="Could not extract text from any provided material files.")

    logger.info("--- [AI Quiz Gen] Context extracted. Generating prompt... ---")
    
    # --- CHANGE 2: UPDATE THE PROMPT ---
    prompt = ChatPromptTemplate.from_messages([
        ("system",
         "You are a quiz generation expert. Create exactly {num_questions} questions "
         "based *only* on the provided context material. The questions should match {difficulty} difficulty. "
         "Include a mix of multiple-choice (one correct answer) and multiple-response (select all that apply) questions. "
         "Assign 1 point to easy, 2 to medium, 3 to hard. Provide 4 options for each question. "
         "**Crucially, for each question, you MUST provide a 'concept_tags' list containing 1 to 3 core concepts or topics from the material that the question is testing.**"),
        ("human",
         "--- CONTEXT MATERIAL ---\n{context}\n\n"
         "Please generate the quiz based *only* on the context above in the required JSON format.")
    ])

    chain = prompt | ai_quiz_gen

    try:
        logger.info("--- [AI Quiz Gen] Invoking AI model... ---")
        result = chain.invoke({
            "num_questions": request.num_questions,
            "difficulty": request.difficulty,
            "context": context_text[:10000]
        })
        logger.info("--- [AI Quiz Gen] AI model returned. Converting questions... ---")
        
        questions_converted = []
        for q in result.questions:
            questions_converted.append(
                schemas.QuestionCreate(
                    question_text=q.question_text,
                    question_type=q.question_type,
                    score=q.score,
                    options=[
                        schemas.OptionCreate(
                            option_text=o.option_text,
                            is_correct=o.is_correct
                        ) for o in q.options
                    ],
                    # --- CHANGE 3: SAVE THE NEW TAGS ---
                    concept_tags=q.concept_tags
                )
            )

        return schemas.QuizGenerationResponse(questions=questions_converted)
    except Exception as e:
        logger.error(f"AI Quiz Gen failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"AI failed to generate quiz: {e}")

# ============================================================
# Get Coursework List
# ============================================================

@router.get(
    "/classrooms/{classroom_id}",
    response_model=List[schemas.CourseworkForStudentList] 
)
def get_courseworks(
    classroom_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_active_user)
):
    student_id_param = None
    if current_user.role == "student":
        if not crud.is_student_enrolled(db, current_user.id, classroom_id):
            raise HTTPException(status_code=403, detail="You are not enrolled in this class")
        student_id_param = current_user.id # <-- Pass student_id
    elif current_user.role == "teacher":
        db_classroom = crud.get_classroom_by_id(db, classroom_id=classroom_id)
        if not db_classroom or db_classroom.teacher_id != current_user.id:
            raise HTTPException(status_code=403, detail="You do not own this classroom")

    return crud.get_courseworks_for_classroom(
        db=db, 
        classroom_id=classroom_id, 
        student_id=student_id_param
    )


# ============================================================
# Get Coursework to Take (with timezone fix)
# ============================================================

@router.get(
    "/{coursework_id}",
    response_model=schemas.CourseworkForStudent
)
def get_coursework_to_take(
    coursework_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_student_user)
):
    db_coursework = crud.get_coursework_with_details(db=db, coursework_id=coursework_id)
    if not db_coursework:
        raise HTTPException(status_code=404, detail="Coursework not found")

    if not crud.is_student_enrolled(db, current_user.id, db_coursework.classroom_id):
        raise HTTPException(status_code=403, detail="You are not enrolled in this coursework's class")

    now = datetime.now(timezone.utc)
    available_from = (
        db_coursework.available_from.replace(tzinfo=timezone.utc)
        if db_coursework.available_from and db_coursework.available_from.tzinfo is None
        else db_coursework.available_from
    )
    due_at = (
        db_coursework.due_at.replace(tzinfo=timezone.utc)
        if db_coursework.due_at and db_coursework.due_at.tzinfo is None
        else db_coursework.due_at
    )
    if available_from and available_from > now:
        raise HTTPException(
            status_code=403,
            detail=f"This coursework is not available until {db_coursework.available_from}"
        )
    if due_at and now > due_at:
        raise HTTPException(status_code=403, detail="The deadline for this coursework has passed")

    existing_submission = crud.get_submission_by_student_and_coursework(db, current_user.id, coursework_id)
    if existing_submission:
        raise HTTPException(status_code=409, detail="You have already submitted this coursework.", headers={"X-Submission-ID": str(existing_submission.id)})

    if db_coursework.rubric_file_url:
        try:
            db_coursework.rubric_file_url = get_presigned_url_for_key(
                db_coursework.rubric_file_url
            )
        except Exception as e:
            logger.error(f"Failed to generate URL for rubric key: {e}")
            db_coursework.rubric_file_url = None
    if db_coursework.material_file_urls:
        fresh_urls = []
        for key in db_coursework.material_file_urls:
            try:
                fresh_urls.append(get_presigned_url_for_key(key))
            except Exception:
                pass 
        db_coursework.material_file_urls = fresh_urls
    return db_coursework

@router.get("/{coursework_id}/details", response_model=schemas.CourseworkDisplay)
def get_coursework_details(
    coursework_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_teacher_user)
):
    """Teacher-only endpoint to get basic coursework info."""
    db_coursework = db.query(models.Coursework).get(coursework_id)
    
    if not db_coursework:
        raise HTTPException(status_code=404, detail="Coursework not found")
    if db_coursework.classroom.teacher_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
        
    return db_coursework

# ============================================================
# Submit Quiz
# ============================================================

@router.post(
    "/{coursework_id}/submit-quiz",
    response_model=schemas.SubmissionDetail
)
def submit_quiz(
    coursework_id: int,
    submission: schemas.QuizSubmissionCreate,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_student_user)
):
    db_coursework = crud.get_coursework_with_details(db=db, coursework_id=coursework_id)
    if not db_coursework or db_coursework.coursework_type != 'quiz':
        raise HTTPException(status_code=404, detail="Quiz coursework not found")
    if not crud.is_student_enrolled(db, current_user.id, db_coursework.classroom_id):
        raise HTTPException(status_code=403, detail="You are not enrolled in this class")

    now = datetime.now(timezone.utc)
    due_at = (
        db_coursework.due_at.replace(tzinfo=timezone.utc)
        if db_coursework.due_at and db_coursework.due_at.tzinfo is None
        else db_coursework.due_at
    )
    if due_at and now > due_at:
        raise HTTPException(status_code=403, detail="The deadline has passed")

    existing_submission = crud.get_submission_by_student_and_coursework(db, current_user.id, coursework_id)
    if existing_submission:
        raise HTTPException(status_code=403, detail="You have already submitted this quiz.")

    db_submission = crud.create_quiz_submission(
        db=db, submission=submission, coursework_id=coursework_id, student_id=current_user.id
    )
    tasks.run_quiz_grading.delay(db_submission.id)

    return crud.get_submission_detail(db, db_submission.id)


# ============================================================
# Submit Essay / Assignment / Case Study
# ============================================================
def handle_file_upload_from_bytes(
    file_obj: io.BytesIO, 
    filename: str, 
    content_type: str, 
    folder: str = "uploads"
) -> str:
    """Helper to upload a file-like object (BytesIO) to MinIO and return its KEY."""
    try:
        file_extension = filename.split(".")[-1]
        unique_filename = f"{folder}/{uuid.uuid4()}.{file_extension}"

        file_key = upload_file_to_storage(
            file_obj=file_obj,
            file_name=unique_filename,
            content_type=content_type
        )
        return file_key
    except Exception as e:
        logger.error(f"File upload from bytes failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"File upload failed: {e}")
    

@router.post(
    "/{coursework_id}/submit-file",
    response_model=schemas.SubmissionDetail
)
async def submit_file(
    coursework_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_student_user)
):
    db_coursework = crud.get_coursework_with_details(db=db, coursework_id=coursework_id)
    if not db_coursework or db_coursework.coursework_type not in ['essay', 'assignment', 'case_study']:
        raise HTTPException(status_code=404, detail="Coursework not found")
    if not crud.is_student_enrolled(db, current_user.id, db_coursework.classroom_id):
        raise HTTPException(status_code=403, detail="You are not enrolled in this class")

    now = datetime.now(timezone.utc)
    due_at = (
        db_coursework.due_at.replace(tzinfo=timezone.utc)
        if db_coursework.due_at and db_coursework.due_at.tzinfo is None
        else db_coursework.due_at
    )
    if due_at and now > due_at:
        raise HTTPException(status_code=403, detail="The deadline has passed")

    existing_submission = crud.get_submission_by_student_and_coursework(db, current_user.id, coursework_id)
    if existing_submission:
        raise HTTPException(status_code=403, detail="You have already submitted this coursework.")
    
    try:
        file_contents = await file.read()
    except Exception as e:
        logger.error(f"Failed to read file stream: {e}")
        raise HTTPException(status_code=500, detail="Failed to read file.")
    
    text_content = ""
    filename = file.filename.lower()
    try:
        if filename.endswith(".pdf"):
            reader = pypdf.PdfReader(io.BytesIO(file_contents))
            text_content = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(file_contents))
            text_content = "\n".join(para.text for para in doc.paragraphs)
        elif filename.endswith(".txt"):
            text_content = file_contents.decode("utf-8")
    except Exception as e:
        logger.error(f"Failed to extract text during submission: {e}")
    
    file_key = handle_file_upload_from_bytes(
        io.BytesIO(file_contents), 
        file.filename, 
        file.content_type
    )

    submission_schema = schemas.EssaySubmissionCreate(
        submission_text=text_content,
        submission_file_url=file_key
    )

    db_submission = crud.create_essay_submission(
        db=db, submission=submission_schema, coursework_id=coursework_id, student_id=current_user.id
    )
    tasks.run_ai_evaluation.delay(db_submission.id)
    return crud.get_submission_detail(db, db_submission.id)


# ============================================================
# Get Submission Result
# ============================================================

@router.get(
    "/submissions/{submission_id}/result",
    response_model=schemas.SubmissionDetail
)
def get_submission_result(
    submission_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_active_user)
):
    db_submission = crud.get_submission_detail(db, submission_id)
    if not db_submission:
        raise HTTPException(status_code=404, detail="Submission not found")
        
    if current_user.role == 'student' and db_submission.student_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
    elif current_user.role == 'teacher' and db_submission.coursework.classroom.teacher_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    if db_submission and db_submission.submission_file_url:
        try:
            db_submission.submission_file_url = get_presigned_url_for_key(
                db_submission.submission_file_url
            )
        except Exception as e:
            logger.error(f"Failed to generate URL for submission key: {e}")
            db_submission.submission_file_url = None
       
    if current_user.role == 'student' and db_submission.status != 'GRADED':
        coursework_data = schemas.CourseworkDisplay.model_validate(db_submission.coursework)
        student_data = schemas.UserDisplay.model_validate(db_submission.student)
        
        return schemas.SubmissionDetail(
            id=db_submission.id, 
            submitted_at=db_submission.submitted_at, 
            status=db_submission.status, 
            coursework_id=db_submission.coursework_id,
            student_id=db_submission.student_id, 
            coursework=coursework_data,
            student=student_data,
            submission_file_url=db_submission.submission_file_url
        )
        
    return db_submission

@router.get("/{coursework_id}/submissions", response_model=List[schemas.SubmissionDetail])
def get_all_submissions(
    coursework_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_teacher_user)
):
    db_coursework = crud.get_coursework_with_details(db, coursework_id)
    if not db_coursework or db_coursework.classroom.teacher_id != current_user.id:
        raise HTTPException(status_code=403, detail="You do not own this coursework")
    
    return crud.get_submissions_for_coursework(db, coursework_id)

@router.patch("/submissions/{submission_id}/approve", response_model=schemas.SubmissionDetail)
def approve_submission(
    submission_id: int,
    approval_data: schemas.SubmissionApproval,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_teacher_user)
):
    db_submission = crud.get_submission_detail(db, submission_id)
    if not db_submission:
        raise HTTPException(status_code=404, detail="Submission not found")
    if db_submission.coursework.classroom.teacher_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
        
    crud.approve_submission(db, submission_id, approval_data)
    tasks.run_dskg_update.delay(submission_id)
    return crud.get_submission_detail(db, submission_id)

@router.delete("/{coursework_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_coursework(
    coursework_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_teacher_user)
):
    db_coursework = crud.get_coursework_with_details(db, coursework_id)
    if not db_coursework:
        raise HTTPException(status_code=404, detail="Coursework not found")

    # Verify the teacher owns the classroom this coursework is in
    if db_coursework.classroom.teacher_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
        
    crud.delete_coursework(db, db_coursework)
    return {"message": "Coursework deleted successfully"}

@router.patch("/questions/{question_id}/options/{option_id}", status_code=status.HTTP_204_NO_CONTENT)
def update_option(
    question_id: int,
    option_id: int,
    approval_data: OptionUpdate, 
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_teacher_user)
):
    db_question = crud.get_question_with_options(db, question_id)
    if not db_question: raise HTTPException(status_code=404)
    if db_question.coursework.classroom.teacher_id != current_user.id: raise HTTPException(status_code=403)
    
    crud.update_option_correctness(db, option_id, approval_data.is_correct)
    
    tasks.regrade_quiz_submissions_for_question.delay(question_id)
    
    return {"message": "Option updated and regrading triggered."}

@router.delete("/{coursework_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_coursework(
    coursework_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_teacher_user)
):
    db_coursework = crud.get_coursework_with_details(db, coursework_id)
    if not db_coursework:
        raise HTTPException(status_code=404, detail="Coursework not found")
    if db_coursework.classroom.teacher_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
        
    crud.delete_coursework(db, db_coursework)
    return {"message": "Coursework deleted successfully"}
# ============================================================
# Rubric Upload + Parse
# ============================================================

@router.post("/upload-rubric", response_model=schemas.RubricParseRequest)
async def upload_rubric_file(
    file: UploadFile = File(...),
    current_user: models.User = Depends(auth.get_teacher_user)
):
    """Extract text from PDF/DOCX/TXT rubric."""
    raw_text = ""
    try:
        contents = await file.read()
        ext = file.filename.split(".")[-1].lower()

        if ext == "txt":
            raw_text = contents.decode("utf-8")
        elif ext == "pdf":
            with io.BytesIO(contents) as f:
                reader = pypdf.PdfReader(f)
                raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext == "docx":
            with io.BytesIO(contents) as f:
                doc = docx.Document(f)
                raw_text = "\n".join(p.text for p in doc.paragraphs)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Upload PDF, DOCX, or TXT.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {e}")

    return schemas.RubricParseRequest(raw_text=raw_text)


@router.post("/parse-rubric", response_model=schemas.RubricParseResponse)
def parse_rubric_with_ai(
    request: schemas.RubricParseRequest,
    current_user: models.User = Depends(auth.get_teacher_user)
):
    """Uses an LLM to convert raw rubric text into structured JSON."""
    structured_rubric_parser = llm.with_structured_output(ParsedRubric)

    prompt = ChatPromptTemplate.from_messages([
        ("system",
         "You are an assistant that parses unstructured text into a JSON rubric. "
         "Identify each grading criterion and its maximum points. "
         "Example: 'Clarity (10 pts) and Grammar (5 pts)' should become "
         "[{{\"criterion\": \"Clarity\", \"max_points\": 10}}, "
         "{{\"criterion\": \"Grammar\", \"max_points\": 5}}]"),
        ("human", "Please parse the following rubric text:\n{raw_text}")
    ])

    chain = prompt | structured_rubric_parser

    try:
        result = chain.invoke({"raw_text": request.raw_text})
        return schemas.RubricParseResponse(rubric=result.rubric)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI failed to parse rubric: {e}")
